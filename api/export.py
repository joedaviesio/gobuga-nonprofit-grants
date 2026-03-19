"""Export draft in multiple formats, including filled templates (multi-tenant)."""

import json
import os
import re
import shutil
import subprocess
from datetime import datetime, timezone
from copy import deepcopy

from api.case_manager import load_case, _case_dir, _uploads_dir
from api.tenant import org_evidence_dir


def _get_case_evidence(org_id: str, case: dict) -> list[dict]:
    """Load evidence items related to this case from the cycle evidence archive."""
    grant_brief = case.get("grant_brief", {})
    source_cycle = grant_brief.get("source_cycle")
    if not source_cycle:
        return []

    # Evidence is now stored in the cycle output dir
    from api.tenant import org_cycles_dir
    evidence_path = os.path.join(org_cycles_dir(org_id), source_cycle, "latest", "evidence.json")
    if not os.path.exists(evidence_path):
        return []

    with open(evidence_path) as f:
        all_evidence = json.load(f)

    if not all_evidence:
        return []

    related_ids = {
        ev.get("id") for ev in grant_brief.get("related_evidence", []) if ev.get("id")
    }

    opp_title = grant_brief.get("title", "").lower()
    title_words = set(opp_title.split()) - {"the", "of", "and", "for", "in", "a", "to"}

    matched = []
    seen_ids = set()
    for ev in all_evidence:
        ev_id = ev.get("id", "")
        if ev_id in seen_ids:
            continue

        if ev_id in related_ids:
            matched.append(ev)
            seen_ids.add(ev_id)
            continue

        if ev.get("type") in ("grant_opportunity", "analysis", "recommendation"):
            ev_title = ev.get("title", "").lower()
            ev_words = set(ev_title.split()) - {"the", "of", "and", "for", "in", "a", "to"}
            if len(title_words & ev_words) >= 2:
                matched.append(ev)
                seen_ids.add(ev_id)

    return matched


def _evidence_appendix_md(evidence: list[dict], source_cycle: str) -> str:
    """Format evidence items as a markdown appendix."""
    if not evidence:
        return ""

    parts = ["\n---\n\n## Sources & Evidence\n"]
    parts.append(f"*From grant cycle: {source_cycle}*\n")

    by_type: dict[str, list[dict]] = {}
    for ev in evidence:
        t = ev.get("type", "other")
        by_type.setdefault(t, []).append(ev)

    type_labels = {
        "grant_opportunity": "Grant Opportunities",
        "analysis": "Analysis",
        "recommendation": "Recommendations",
        "donor_intel": "Donor Intelligence",
        "deadline": "Deadlines",
    }

    for ev_type, items in by_type.items():
        label = type_labels.get(ev_type, ev_type.replace("_", " ").title())
        parts.append(f"### {label}\n")
        for ev in items:
            title = ev.get("title", "Untitled")
            parts.append(f"**{title}** ({ev.get('id', '')})")
            if ev.get("source_url"):
                parts.append(f"  Source: {ev['source_url']}")
            content = ev.get("content", "")
            if content:
                if len(content) > 500:
                    content = content[:500] + "..."
                parts.append(f"  {content}")
            if ev.get("severity"):
                parts.append(f"  Severity: {ev['severity']}")
            parts.append("")

    return "\n".join(parts)


def export_markdown(org_id: str, case_id: str) -> str | None:
    """Export the current draft as a markdown document."""
    case = load_case(org_id, case_id)
    if case is None:
        return None

    sections = case.get("draft", {}).get("sections", {})
    if not sections:
        return "# Draft\n\nNo sections drafted yet."

    parts = [f"# Grant Application: {case.get('grant_id', 'Unknown')}\n"]
    parts.append(f"**Case:** {case['case_id']}")
    parts.append(f"**Status:** {case['status']}")
    parts.append(f"**Generated:** {datetime.now(timezone.utc).isoformat()}\n")
    parts.append("---\n")

    for name, sec in sections.items():
        title = name.replace("_", " ").title()
        status = sec.get("status", "draft")
        content = sec.get("content", "")
        parts.append(f"## {title}\n")
        if status != "complete":
            parts.append(f"*Status: {status}*\n")
        parts.append(content)
        parts.append("")

    evidence = _get_case_evidence(org_id, case)
    source_cycle = case.get("grant_brief", {}).get("source_cycle", "")
    appendix = _evidence_appendix_md(evidence, source_cycle)
    if appendix:
        parts.append(appendix)

    return "\n".join(parts)


def export_json(org_id: str, case_id: str) -> dict | None:
    """Export the current draft as structured JSON."""
    case = load_case(org_id, case_id)
    if case is None:
        return None

    evidence = _get_case_evidence(org_id, case)
    return {
        "case_id": case["case_id"],
        "grant_id": case.get("grant_id"),
        "status": case["status"],
        "exported": datetime.now(timezone.utc).isoformat(),
        "sections": case.get("draft", {}).get("sections", {}),
        "evidence": [
            {
                "id": ev.get("id"),
                "type": ev.get("type"),
                "title": ev.get("title"),
                "source_url": ev.get("source_url"),
                "severity": ev.get("severity"),
                "content": ev.get("content", "")[:500],
            }
            for ev in evidence
        ],
    }


def _md_to_plain(text: str) -> str:
    """Strip markdown formatting to plain text for DOCX insertion."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^[-*]\s+', '- ', text, flags=re.MULTILINE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _doc_to_docx(doc_path: str) -> str | None:
    """Convert a legacy .doc to .docx using textutil (macOS)."""
    if not shutil.which("textutil"):
        return None
    docx_path = doc_path.rsplit(".", 1)[0] + ".docx"
    result = subprocess.run(
        ["textutil", "-convert", "docx", "-output", docx_path, doc_path],
        capture_output=True, text=True, timeout=30,
    )
    if result.returncode == 0 and os.path.exists(docx_path):
        return docx_path
    return None


def export_docx(org_id: str, case_id: str, template_filename: str | None = None) -> str | None:
    """Export to DOCX. Returns the filepath of the generated DOCX."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    case = load_case(org_id, case_id)
    if case is None:
        return None

    sections = case.get("draft", {}).get("sections", {})
    case_dir = _case_dir(org_id, case_id)
    exports_dir = os.path.join(case_dir, "exports")
    os.makedirs(exports_dir, exist_ok=True)
    ts = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")

    if template_filename:
        # Convert .doc to .docx first if needed
        if template_filename.lower().endswith(".doc") and not template_filename.lower().endswith(".docx"):
            uploads_dir = _uploads_dir(org_id, case_id)
            doc_path = os.path.join(uploads_dir, template_filename)
            converted = _doc_to_docx(doc_path)
            if converted:
                template_filename = os.path.basename(converted)
            else:
                # Can't convert .doc — fall back to fresh DOCX generation
                template_filename = None

    if template_filename:
        filepath = _fill_docx_template(org_id, case_id, template_filename, sections, exports_dir, ts)
    else:
        filepath = _generate_fresh_docx(org_id, case_id, sections, exports_dir, ts, case)

    if filepath:
        case["exports"].append({
            "format": "docx",
            "filename": os.path.basename(filepath),
            "timestamp": datetime.now(timezone.utc).isoformat(),
        })
        from api.case_manager import _save
        _save(org_id, case)

    return filepath


def _build_field_map(org_id: str, case_id: str, sections: dict) -> dict:
    """Build a field label -> content mapping from draft sections and parsed metadata.

    Uses the parsed_sections labels (from Bot B) as keys so they match the
    original template field names, and maps them to the filled content (from Bot C).
    Also indexes by section_id for broader matching.
    """
    field_map = {}

    # Get parsed section metadata for label mapping
    case = load_case(org_id, case_id)
    parsed = (case.get("parsed_sections") or {}).get("sections", []) if case else []
    label_by_id = {s["id"]: s.get("label", s["id"]) for s in parsed}

    for section_id, sec in sections.items():
        content = sec.get("content", "").strip()
        if not content:
            continue
        plain = _md_to_plain(content)

        # Key by the original label from the parsed template
        label = label_by_id.get(section_id, section_id.replace("_", " "))
        field_map[label.lower()] = plain

        # Also key by section_id variants for broader matching
        field_map[section_id.lower()] = plain
        field_map[section_id.replace("_", " ").lower()] = plain

    return field_map


def _fill_docx_template(org_id: str, case_id: str, template_filename: str, sections: dict, exports_dir: str, ts: str) -> str | None:
    """Fill an uploaded DOCX template with draft content."""
    from docx import Document

    template_path = os.path.join(_uploads_dir(org_id, case_id), template_filename)
    if not os.path.exists(template_path):
        return None

    doc = Document(template_path)

    # Build field map directly from draft sections keyed by section_id and label
    field_map = _build_field_map(org_id, case_id, sections)
    _fill_document(doc, field_map)

    output_filename = f"filled-{template_filename.replace('.docx', '')}-{ts}.docx"
    output_path = os.path.join(exports_dir, output_filename)
    doc.save(output_path)
    return output_path


def _parse_form_content(content: str) -> dict:
    """Parse the bot's filled form content into a field->value mapping."""
    field_map = {}
    lines = content.split("\n")
    current_field = None
    current_value_lines = []
    current_section = None
    section_lines = []

    for line in lines:
        section_match = re.match(r'^###\s+(.+)', line)
        if section_match:
            if current_section:
                section_text = _md_to_plain("\n".join(section_lines).strip())
                if section_text:
                    field_map[current_section] = section_text
            current_section = section_match.group(1).strip().lower()
            section_lines = []
            if current_field:
                field_map[current_field] = _md_to_plain("\n".join(current_value_lines).strip())
                current_field = None
                current_value_lines = []
            continue

        field_match = re.match(r'\*\*(.+?):\*\*\s*(.*)', line)
        if field_match:
            if current_field:
                field_map[current_field] = _md_to_plain("\n".join(current_value_lines).strip())
            current_field = field_match.group(1).strip().lower()
            inline_val = field_match.group(2).strip()
            current_value_lines = [inline_val] if inline_val else []
        elif line.startswith("## ") or line.startswith("---"):
            if current_field:
                field_map[current_field] = _md_to_plain("\n".join(current_value_lines).strip())
                current_field = None
                current_value_lines = []
            if current_section:
                section_text = _md_to_plain("\n".join(section_lines).strip())
                if section_text:
                    field_map[current_section] = section_text
                current_section = None
                section_lines = []
        else:
            if current_field is not None:
                current_value_lines.append(line)
            if current_section is not None:
                section_lines.append(line)

    if current_field:
        field_map[current_field] = _md_to_plain("\n".join(current_value_lines).strip())
    if current_section:
        section_text = _md_to_plain("\n".join(section_lines).strip())
        if section_text:
            field_map[current_section] = section_text

    return field_map


def _normalize(text: str) -> str:
    return re.sub(r'[^a-z0-9\s]', '', text.lower()).strip()


class FieldMapper:
    """Cached field matcher that avoids O(n^2) lookups.

    On first construction, builds a normalized-key lookup dict for O(1) exact
    matches. Substring and fuzzy matches are tried only on cache miss and
    their results are memoized so the same label is never re-scanned.
    """

    def __init__(self, field_map: dict):
        self._field_map = field_map
        # Pre-compute normalized keys once
        self._norm_keys: list[tuple[str, str, set[str]]] = []
        # exact normalized-key -> value
        self._exact: dict[str, str] = {}
        for key, val in field_map.items():
            nk = _normalize(key)
            if nk:
                self._exact[nk] = val
                self._norm_keys.append((nk, val, set(nk.split())))
        # Memoization cache: queried normalized label -> value | None
        self._cache: dict[str, str | None] = {}

    def get_value(self, label: str) -> str | None:
        """Return the best matching value for *label*, or None."""
        norm_label = _normalize(label)
        if not norm_label:
            return None

        # Check memo cache first
        if norm_label in self._cache:
            return self._cache[norm_label]

        result = self._lookup(norm_label)
        self._cache[norm_label] = result
        return result

    def _lookup(self, norm_label: str) -> str | None:
        # 1. Exact match (O(1) dict lookup)
        if norm_label in self._exact:
            return self._exact[norm_label]

        # 2. Substring containment
        for nk, val, _ in self._norm_keys:
            if nk in norm_label or norm_label in nk:
                return val

        # 3. Fuzzy word-overlap (>= 50%)
        label_words = set(norm_label.split())
        best_score = 0.0
        best_val = None
        for nk, val, key_words in self._norm_keys:
            if not key_words:
                continue
            overlap = len(label_words & key_words) / max(len(label_words), len(key_words))
            if overlap > best_score and overlap >= 0.5:
                best_score = overlap
                best_val = val
        return best_val


def _find_best_match(label: str, field_map: dict) -> str | None:
    """Legacy helper — wraps FieldMapper for one-off calls."""
    mapper = FieldMapper(field_map)
    return mapper.get_value(label)


def _set_cell_text(cell, text: str):
    """Set cell text while preserving the first paragraph's formatting."""
    # Clear all paragraphs except the first
    for p in cell.paragraphs[1:]:
        p_element = p._element
        p_element.getparent().remove(p_element)
    # Set text on the first paragraph, preserving its style
    if cell.paragraphs:
        p = cell.paragraphs[0]
        # Preserve the paragraph format and style
        if p.runs:
            # Keep the first run's formatting, clear the rest
            first_run = p.runs[0]
            for run in p.runs[1:]:
                run._element.getparent().remove(run._element)
            first_run.text = text
        else:
            p.text = text


def _set_para_text(para, text: str):
    """Set paragraph text while preserving formatting from existing runs."""
    if para.runs:
        first_run = para.runs[0]
        for run in para.runs[1:]:
            run._element.getparent().remove(run._element)
        first_run.text = text
    else:
        para.text = text


def _fill_document(doc, field_map: dict):
    """Walk document tables and paragraphs, filling cells that match field labels."""
    mapper = FieldMapper(field_map)
    filled_keys = set()

    for table in doc.tables:
        for ri, row in enumerate(table.rows):
            cells = row.cells
            for i, cell in enumerate(cells):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue

                match = mapper.get_value(cell_text)
                if match and match != cell_text:
                    # Try filling adjacent cell (label | value layout)
                    if i + 1 < len(cells):
                        target = cells[i + 1]
                        if not target.text.strip() or target.text.strip() == cell_text:
                            _set_cell_text(target, match)
                            filled_keys.add(_normalize(cell_text))
                            continue

                    # Try filling the row below (label-above-value layout)
                    if ri + 1 < len(table.rows):
                        below_cell = table.rows[ri + 1].cells[i]
                        if not below_cell.text.strip():
                            _set_cell_text(below_cell, match)
                            filled_keys.add(_normalize(cell_text))
                            continue

                    # Fallback: append to same cell if multi-paragraph
                    if len(cell.paragraphs) > 1:
                        _set_para_text(cell.paragraphs[-1], match)
                        filled_keys.add(_normalize(cell_text))

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Fill placeholder markers
        if "[TO BE COMPLETED" in text or "[TBD]" in text or "[INSERT" in text:
            match = mapper.get_value(text)
            if match:
                _set_para_text(para, match)
                continue
        # Fill empty paragraphs that follow a label paragraph
        match = mapper.get_value(text)
        if match and match != text and _normalize(text) not in filled_keys:
            # Check if this looks like a section heading/label (short text, no answer content)
            if len(text) < 200 and not any(c in text for c in ['.', ',', ';']) or text.endswith(':') or text.endswith('?'):
                _set_para_text(para, text + "\n" + match)


def _generate_fresh_docx(org_id: str, case_id: str, sections: dict, exports_dir: str, ts: str, case: dict) -> str:
    """Generate a fresh DOCX document from draft sections."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    title = doc.add_heading(f"Grant Application: {case.get('grant_id', 'Unknown')}", level=0)
    doc.add_paragraph(f"Case: {case['case_id']}")
    doc.add_paragraph(f"Status: {case['status']}")
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")

    for name, sec in sections.items():
        title_text = name.replace("_", " ").title()
        doc.add_heading(title_text, level=1)

        if sec.get("status") and sec["status"] != "complete":
            status_para = doc.add_paragraph(f"Status: {sec['status']}")
            status_para.runs[0].italic = True

        content = sec.get("content", "")
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue

            if line.startswith("### "):
                doc.add_heading(_md_to_plain(line[4:]), level=3)
            elif line.startswith("## "):
                doc.add_heading(_md_to_plain(line[3:]), level=2)
            elif line.startswith("# "):
                doc.add_heading(_md_to_plain(line[2:]), level=1)
            elif line.startswith("- "):
                doc.add_paragraph(_md_to_plain(line[2:]), style="List Bullet")
            else:
                doc.add_paragraph(_md_to_plain(line))

    evidence = _get_case_evidence(org_id, case)
    if evidence:
        source_cycle = case.get("grant_brief", {}).get("source_cycle", "")
        doc.add_page_break()
        doc.add_heading("Sources & Evidence", level=1)
        doc.add_paragraph(f"From grant cycle: {source_cycle}")

        type_labels = {
            "grant_opportunity": "Grant Opportunities",
            "analysis": "Analysis",
            "recommendation": "Recommendations",
            "donor_intel": "Donor Intelligence",
            "deadline": "Deadlines",
        }
        by_type: dict[str, list[dict]] = {}
        for ev in evidence:
            t = ev.get("type", "other")
            by_type.setdefault(t, []).append(ev)

        for ev_type, items in by_type.items():
            label = type_labels.get(ev_type, ev_type.replace("_", " ").title())
            doc.add_heading(label, level=2)
            for ev in items:
                title = ev.get("title", "Untitled")
                doc.add_paragraph(f"{title} ({ev.get('id', '')})", style="List Bullet")
                if ev.get("source_url"):
                    doc.add_paragraph(f"Source: {ev['source_url']}")
                content = ev.get("content", "")
                if content:
                    if len(content) > 500:
                        content = content[:500] + "..."
                    doc.add_paragraph(content)

    output_filename = f"draft-{ts}.docx"
    output_path = os.path.join(exports_dir, output_filename)
    doc.save(output_path)
    return output_path


def _fill_pdf_template(org_id: str, case_id: str, template_filename: str, sections: dict, exports_dir: str, ts: str) -> str | None:
    """Fill a PDF template using form fields (AcroForm) or text overlay near labels."""
    import fitz  # PyMuPDF

    template_path = os.path.join(_uploads_dir(org_id, case_id), template_filename)
    if not os.path.exists(template_path):
        return None

    field_map = _build_field_map(org_id, case_id, sections)
    mapper = FieldMapper(field_map)
    doc = fitz.open(template_path)

    # Strategy 1: Fill AcroForm fields if the PDF has them
    form_filled = False
    for page in doc:
        widgets = list(page.widgets())
        for widget in widgets:
            field_name = widget.field_name or ""
            if not field_name:
                continue
            match = mapper.get_value(field_name)
            if match:
                widget.field_value = match
                widget.update()
                form_filled = True

    # Strategy 2: If no form fields, find label text and insert content below
    if not form_filled:
        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("type") != 0:  # text block
                    continue
                for line in block.get("lines", []):
                    line_text = "".join(span["text"] for span in line["spans"]).strip()
                    if not line_text or len(line_text) > 200:
                        continue
                    match = mapper.get_value(line_text)
                    if match and match != line_text:
                        # Insert below the label line
                        bbox = line["bbox"]
                        insert_y = bbox[3] + 2
                        # Use a reasonable font size
                        font_size = 9
                        if line["spans"]:
                            font_size = min(line["spans"][0].get("size", 9), 10)
                        # Truncate to fit page width
                        max_width = page.rect.width - bbox[0] - 20
                        text_rect = fitz.Rect(bbox[0], insert_y, bbox[0] + max_width, insert_y + 200)
                        page.insert_textbox(
                            text_rect,
                            match,
                            fontsize=font_size,
                            fontname="helv",
                            color=(0, 0, 0),
                        )

    stem = template_filename.rsplit(".", 1)[0]
    output_filename = f"filled-{stem}-{ts}.pdf"
    output_path = os.path.join(exports_dir, output_filename)
    doc.save(output_path)
    doc.close()
    return output_path


def export_pdf(org_id: str, case_id: str, template_filename: str | None = None) -> str | None:
    """Export the current draft as a PDF document. If template_filename is a PDF, fill it."""
    case = load_case(org_id, case_id)
    if case is None:
        return None

    sections = case.get("draft", {}).get("sections", {})
    case_dir = _case_dir(org_id, case_id)
    exports_dir = os.path.join(case_dir, "exports")
    os.makedirs(exports_dir, exist_ok=True)
    ts = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")

    # If a PDF template was uploaded, fill it
    if template_filename and template_filename.lower().endswith(".pdf"):
        filepath = _fill_pdf_template(org_id, case_id, template_filename, sections, exports_dir, ts)
        if filepath:
            case["exports"].append({
                "format": "pdf",
                "filename": os.path.basename(filepath),
                "timestamp": datetime.now(timezone.utc).isoformat(),
            })
            from api.case_manager import _save
            _save(org_id, case)
            return filepath

    from fpdf import FPDF
    ts = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")

    def _pdf_safe(text: str) -> str:
        """Replace Unicode characters that Helvetica can't render."""
        replacements = {
            "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
            "\u2013": "-", "\u2014": "--", "\u2026": "...", "\u00a0": " ",
            "\u2022": "-", "\u2023": ">", "\u25cf": "-", "\u2010": "-",
            "\u2011": "-", "\u2012": "-", "\u200b": "", "\u200c": "",
            "\u200d": "", "\ufeff": "",
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text.encode("latin-1", errors="replace").decode("latin-1")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, _pdf_safe(f"Grant Application: {case.get('grant_id', 'Unknown')}"), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, _pdf_safe(f"Case: {case['case_id']}"), new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, _pdf_safe(f"Status: {case['status']}"), new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(8)

    for name, sec in sections.items():
        title_text = name.replace("_", " ").title()
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 8, _pdf_safe(title_text), new_x="LMARGIN", new_y="NEXT")

        if sec.get("status") and sec["status"] != "complete":
            pdf.set_font("Helvetica", "I", 9)
            pdf.cell(0, 5, _pdf_safe(f"Status: {sec['status']}"), new_x="LMARGIN", new_y="NEXT")

        content = _md_to_plain(sec.get("content", ""))
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, _pdf_safe(content))
        pdf.ln(4)

    evidence = _get_case_evidence(org_id, case)
    if evidence:
        source_cycle = case.get("grant_brief", {}).get("source_cycle", "")
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, "Sources & Evidence", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(0, 5, _pdf_safe(f"From grant cycle: {source_cycle}"), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)
        for ev in evidence:
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 6, _pdf_safe(f"{ev.get('title', 'Untitled')} ({ev.get('id', '')})"), new_x="LMARGIN", new_y="NEXT")
            if ev.get("source_url"):
                pdf.set_font("Helvetica", "", 8)
                pdf.cell(0, 4, _pdf_safe(f"Source: {ev['source_url']}"), new_x="LMARGIN", new_y="NEXT")
            content = ev.get("content", "")
            if content:
                if len(content) > 500:
                    content = content[:500] + "..."
                pdf.set_font("Helvetica", "", 9)
                pdf.multi_cell(0, 4, _pdf_safe(content))
            pdf.ln(2)

    output_filename = f"draft-{ts}.pdf"
    output_path = os.path.join(exports_dir, output_filename)
    pdf.output(output_path)

    case["exports"].append({
        "format": "pdf",
        "filename": output_filename,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    })
    from api.case_manager import _save
    _save(org_id, case)
    return output_path


def export_xlsx(org_id: str, case_id: str, template_filename: str | None = None) -> str | None:
    """Export to XLSX. Template-fill or fresh generation."""
    from openpyxl import load_workbook, Workbook

    case = load_case(org_id, case_id)
    if case is None:
        return None

    sections = case.get("draft", {}).get("sections", {})
    case_dir = _case_dir(org_id, case_id)
    exports_dir = os.path.join(case_dir, "exports")
    os.makedirs(exports_dir, exist_ok=True)
    ts = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")

    if template_filename:
        filepath = _fill_xlsx_template(org_id, case_id, template_filename, sections, exports_dir, ts)
    else:
        filepath = _generate_fresh_xlsx(org_id, case_id, sections, exports_dir, ts, case)

    if filepath:
        case["exports"].append({
            "format": "xlsx",
            "filename": os.path.basename(filepath),
            "timestamp": datetime.now(timezone.utc).isoformat(),
        })
        from api.case_manager import _save
        _save(org_id, case)

    return filepath


def _fill_xlsx_template(org_id: str, case_id: str, template_filename: str, sections: dict, exports_dir: str, ts: str) -> str | None:
    """Fill an uploaded XLSX template with draft content."""
    from openpyxl import load_workbook

    template_path = os.path.join(_uploads_dir(org_id, case_id), template_filename)
    if not os.path.exists(template_path):
        return None

    wb = load_workbook(template_path)

    field_map = _build_field_map(org_id, case_id, sections)
    mapper = FieldMapper(field_map)

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for i, cell in enumerate(row):
                if cell.value and isinstance(cell.value, str) and cell.value.strip():
                    match = mapper.get_value(cell.value.strip())
                    if match and match != cell.value.strip():
                        if i + 1 < len(row):
                            target = row[i + 1]
                            if not target.value or str(target.value).strip() == "":
                                target.value = match

    stem = template_filename.rsplit(".", 1)[0]
    output_filename = f"filled-{stem}-{ts}.xlsx"
    output_path = os.path.join(exports_dir, output_filename)
    wb.save(output_path)
    return output_path


def _generate_fresh_xlsx(org_id: str, case_id: str, sections: dict, exports_dir: str, ts: str, case: dict) -> str:
    """Generate a fresh XLSX from draft sections."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    ws = wb.active
    ws.title = "Application"

    ws.append([f"Grant Application: {case.get('grant_id', 'Unknown')}"])
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([f"Case: {case['case_id']}"])
    ws.append([f"Status: {case['status']}"])
    ws.append([f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"])
    ws.append([])

    for name, sec in sections.items():
        title_text = name.replace("_", " ").title()
        ws.append([title_text])
        ws.cell(row=ws.max_row, column=1).font = Font(bold=True, size=12)
        content = _md_to_plain(sec.get("content", ""))
        for line in content.split("\n"):
            if line.strip():
                ws.append([line.strip()])
        ws.append([])

    ws.column_dimensions["A"].width = 80

    output_filename = f"draft-{ts}.xlsx"
    output_path = os.path.join(exports_dir, output_filename)
    wb.save(output_path)
    return output_path


def _docx_to_doc(docx_path: str) -> str | None:
    """Convert a DOCX to legacy DOC format using textutil (macOS)."""
    if not shutil.which("textutil"):
        return None
    doc_path = docx_path.rsplit(".", 1)[0] + ".doc"
    result = subprocess.run(
        ["textutil", "-convert", "doc", "-output", doc_path, docx_path],
        capture_output=True, text=True, timeout=30,
    )
    if result.returncode == 0 and os.path.exists(doc_path):
        return doc_path
    return None


def export_to_file(org_id: str, case_id: str, fmt: str = "markdown", template_filename: str | None = None) -> str | None:
    """Export and save to the case directory. Returns the file path."""
    case = load_case(org_id, case_id)
    if case is None:
        return None

    case_dir = _case_dir(org_id, case_id)
    exports_dir = os.path.join(case_dir, "exports")
    os.makedirs(exports_dir, exist_ok=True)

    ts = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")

    if fmt == "markdown":
        content = export_markdown(org_id, case_id)
        filename = f"draft-{ts}.md"
        filepath = os.path.join(exports_dir, filename)
        with open(filepath, "w") as f:
            f.write(content)
    elif fmt == "json":
        content = export_json(org_id, case_id)
        filename = f"draft-{ts}.json"
        filepath = os.path.join(exports_dir, filename)
        with open(filepath, "w") as f:
            json.dump(content, f, indent=2)
    elif fmt == "docx":
        return export_docx(org_id, case_id, template_filename)
    elif fmt == "doc":
        docx_path = export_docx(org_id, case_id, template_filename)
        if not docx_path:
            return None
        doc_path = _docx_to_doc(docx_path)
        if doc_path:
            case = load_case(org_id, case_id)
            case["exports"].append({
                "format": "doc",
                "filename": os.path.basename(doc_path),
                "timestamp": datetime.now(timezone.utc).isoformat(),
            })
            from api.case_manager import _save
            _save(org_id, case)
            return doc_path
        return docx_path
    elif fmt == "pdf":
        return export_pdf(org_id, case_id, template_filename)
    elif fmt == "xlsx":
        return export_xlsx(org_id, case_id, template_filename)
    elif fmt == "xls":
        return export_xlsx(org_id, case_id, template_filename)
    else:
        return None

    case["exports"].append({
        "format": fmt,
        "filename": filename,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    })
    from api.case_manager import _save
    _save(org_id, case)

    return filepath
