"""Shared file text extraction module.

Centralises PDF, DOCX, DOC, XLSX, XLS, HTML, CSV, and plain-text extraction
so that every part of the codebase uses the same logic.
"""

import os
import re


# ---------------------------------------------------------------------------
# Individual format extractors
# ---------------------------------------------------------------------------

def extract_pdf(filepath: str, max_pages: int = 50) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    import fitz
    doc = fitz.open(filepath)
    pages = min(len(doc), max_pages)
    text_parts = []
    for i in range(pages):
        page_text = doc[i].get_text()
        if page_text.strip():
            text_parts.append(f"[Page {i+1}]\n{page_text.strip()}")
    if len(doc) > max_pages:
        text_parts.append(f"\n... [{len(doc) - max_pages} more pages not extracted]")
    doc.close()
    return "\n\n".join(text_parts)


def extract_docx(filepath: str) -> str:
    """Extract text from a .docx file (paragraphs + tables)."""
    from docx import Document
    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_doc(filepath: str) -> str:
    """Extract text from a legacy .doc file using textutil (macOS) or antiword."""
    import subprocess
    import shutil
    if shutil.which("textutil"):
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", filepath],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    if shutil.which("antiword"):
        result = subprocess.run(
            ["antiword", filepath],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    raise RuntimeError(
        "Cannot extract .doc file — install antiword or use macOS textutil"
    )


def extract_xlsx(filepath: str) -> str:
    """Extract text from an .xlsx file."""
    from openpyxl import load_workbook
    wb = load_workbook(filepath, read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"[Sheet: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def extract_xls(filepath: str) -> str:
    """Extract text from a legacy .xls file."""
    import xlrd
    wb = xlrd.open_workbook(filepath)
    parts = []
    for sheet in wb.sheets():
        parts.append(f"[Sheet: {sheet.name}]")
        for row_idx in range(sheet.nrows):
            cells = [str(c).strip() for c in sheet.row_values(row_idx) if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_csv(filepath: str) -> str:
    """Read a CSV file as plain text."""
    with open(filepath, errors="replace") as f:
        return f.read()


def extract_html(filepath: str) -> str:
    """Extract readable text from an HTML file."""
    with open(filepath, errors="replace") as f:
        return strip_html(f.read())


def extract_plain(filepath: str) -> str:
    """Read a plain-text file (.txt, .md, .json, etc.)."""
    with open(filepath, errors="replace") as f:
        return f.read()


def strip_html(html: str) -> str:
    """Strip HTML tags and return readable text."""
    text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
    text = re.sub(r'<(?:br|p|div|h[1-6]|li|tr)[^>]*>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n[ \t]+', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

# Extension -> extractor mapping
_EXTRACTORS = {
    ".pdf": lambda fp, **kw: extract_pdf(fp, max_pages=kw.get("max_pages", 50)),
    ".docx": lambda fp, **kw: extract_docx(fp),
    ".doc": lambda fp, **kw: extract_doc(fp),
    ".xlsx": lambda fp, **kw: extract_xlsx(fp),
    ".xls": lambda fp, **kw: extract_xls(fp),
    ".html": lambda fp, **kw: extract_html(fp),
    ".htm": lambda fp, **kw: extract_html(fp),
    ".csv": lambda fp, **kw: extract_csv(fp),
    ".md": lambda fp, **kw: extract_plain(fp),
    ".txt": lambda fp, **kw: extract_plain(fp),
    ".json": lambda fp, **kw: extract_plain(fp),
}


def extract_text(filepath: str, max_pages: int = 50) -> str:
    """Dispatch to the correct extractor based on file extension.

    Returns the extracted text, or an empty string if the format is
    unsupported or an error occurs.
    """
    ext = os.path.splitext(filepath)[1].lower()
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        # Fall back to plain-text for unknown extensions
        try:
            return extract_plain(filepath)
        except Exception:
            return ""
    try:
        return extractor(filepath, max_pages=max_pages)
    except Exception:
        return ""
