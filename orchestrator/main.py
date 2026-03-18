"""Cycle engine — runs the daily loop (multi-tenant)."""

import json
import os
from datetime import date as date_type
from concurrent.futures import ThreadPoolExecutor, as_completed

from orchestrator.config import agents_by_phase, AGENTS
from orchestrator.agent_runner import run_agent
from orchestrator.state_manager import load_all_state, save_state
from orchestrator.evidence import load_evidence_for_date
from orchestrator.signoff import load_signoff

from api.tenant import org_root, org_cycles_dir, org_data_dir, org_prompts_dir


def _load_prompt(org_id: str, agent_config: dict, state: dict, prior_outputs: dict, date: str) -> str:
    """Build the full system prompt for an agent."""
    # Load prompt from org-specific prompts dir, fall back to templates
    prompt_filename = os.path.basename(agent_config["prompt_file"])
    prompt_path = os.path.join(org_prompts_dir(org_id), prompt_filename)
    if not os.path.exists(prompt_path):
        # Fall back to base templates
        project_root = os.path.dirname(os.path.dirname(__file__))
        prompt_path = os.path.join(project_root, "templates", "prompts", prompt_filename)

    with open(prompt_path) as f:
        role_brief = f.read()

    # Load previous signoff if it exists
    signoff = load_signoff(org_id, date)

    parts = [
        role_brief,
        "\n\n---\n## Current State\n```json\n" + json.dumps(state, indent=2, default=str) + "\n```",
    ]

    if signoff:
        parts.append("\n\n---\n## Previous Signoff Decisions\n```json\n" + json.dumps(signoff, indent=2) + "\n```")

    # Include outputs from dependency agents
    deps = agent_config.get("depends_on", [])
    if deps:
        dep_section = "\n\n---\n## Outputs From Prior Phases\n"
        for dep_id in deps:
            if dep_id in prior_outputs:
                dep_section += f"\n### {dep_id}\n```\n{prior_outputs[dep_id].get('raw_text', 'No output')}\n```\n"
        parts.append(dep_section)

    # Include input data files
    data_section = _load_input_data(org_id)
    if data_section:
        parts.append(data_section)

    return "\n".join(parts)


def _extract_docx(filepath: str) -> str:
    """Extract text from a .docx file."""
    from docx import Document
    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(filepath: str, max_pages: int = 20) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    import fitz
    doc = fitz.open(filepath)
    pages = min(len(doc), max_pages)
    text_parts = []
    for i in range(pages):
        page_text = doc[i].get_text()
        if page_text.strip():
            text_parts.append(f"[Page {i+1}]\n{page_text.strip()}")
    if len(doc) > max_pages:
        text_parts.append(f"\n... [{len(doc) - max_pages} more pages not extracted]")
    doc.close()
    return "\n\n".join(text_parts)


def _strip_html(html: str) -> str:
    """Extract readable text from HTML content."""
    import re
    text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
    text = re.sub(r'<(?:br|p|div|h[1-6]|li|tr)[^>]*>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n[ \t]+', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _load_input_data(org_id: str) -> str:
    """Load all text-based and HTML files from org data/ subfolders into a context section."""
    data_dir = org_data_dir(org_id)
    if not os.path.exists(data_dir):
        return ""

    TEXT_EXTENSIONS = {".md", ".txt", ".json", ".csv"}
    HTML_EXTENSIONS = {".html", ".htm"}
    PDF_EXTENSIONS = {".pdf"}
    DOCX_EXTENSIONS = {".docx"}
    SUPPORTED = TEXT_EXTENSIONS | HTML_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS
    MAX_PER_FILE = 8000

    SKIP_DIRS = {"archive"}

    sections = []
    for dirpath, dirnames, filenames in sorted(os.walk(data_dir)):
        dirnames[:] = sorted(d for d in dirnames if d not in SKIP_DIRS)
        rel_dir = os.path.relpath(dirpath, data_dir)
        if rel_dir == ".":
            continue

        files_content = []
        for fname in sorted(filenames):
            ext = os.path.splitext(fname)[1].lower()
            if ext not in SUPPORTED:
                continue
            filepath = os.path.join(dirpath, fname)
            try:
                if ext in PDF_EXTENSIONS:
                    content = _extract_pdf(filepath)
                elif ext in DOCX_EXTENSIONS:
                    content = _extract_docx(filepath)
                elif ext in HTML_EXTENSIONS:
                    with open(filepath, errors="replace") as f:
                        content = _strip_html(f.read())
                else:
                    with open(filepath, errors="replace") as f:
                        content = f.read()
                if len(content) > MAX_PER_FILE:
                    content = content[:MAX_PER_FILE] + f"\n... [truncated, {len(content)} chars total]"
                files_content.append(f"#### {fname}\n```\n{content}\n```")
            except Exception:
                continue

        if files_content:
            sections.append(f"### {rel_dir}/\n" + "\n\n".join(files_content))

    if not sections:
        return ""

    return "\n\n---\n## Input Data\n\n" + "\n\n".join(sections)


def run_cycle(org_id: str, cycle_date: str | None = None, model_overrides: dict | None = None):
    """Run a full cycle for the given org and date."""
    if cycle_date is None:
        cycle_date = date_type.today().isoformat()

    print(f"\n{'='*60}")
    print(f"  Grant Cycle — {cycle_date} — org: {org_id}")
    print(f"{'='*60}\n")

    state = load_all_state(org_id)
    phases = agents_by_phase()
    all_outputs = {}

    # Apply model overrides if provided (e.g. for tier-based model selection)
    agents = []
    for phase_num, phase_agents in phases.items():
        for agent_cfg in phase_agents:
            cfg = dict(agent_cfg)
            if model_overrides and cfg["id"] in model_overrides:
                cfg["model"] = model_overrides[cfg["id"]]
            agents.append((phase_num, cfg))

    # Re-group by phase
    phase_groups = {}
    for phase_num, cfg in agents:
        phase_groups.setdefault(phase_num, []).append(cfg)

    for phase_num in sorted(phase_groups.keys()):
        phase_agents = phase_groups[phase_num]
        print(f"--- Phase {phase_num} ---")

        if phase_num == 1:
            with ThreadPoolExecutor(max_workers=len(phase_agents)) as executor:
                futures = {}
                for agent_cfg in phase_agents:
                    prompt = _load_prompt(org_id, agent_cfg, state, all_outputs, cycle_date)
                    future = executor.submit(run_agent, org_id, agent_cfg, prompt, cycle_date)
                    futures[future] = agent_cfg

                for future in as_completed(futures):
                    agent_cfg = futures[future]
                    try:
                        output = future.result()
                        all_outputs[agent_cfg["id"]] = output
                        print(f"  ✓ {agent_cfg['name']} complete")
                    except Exception as e:
                        print(f"  ✗ {agent_cfg['name']} failed: {e}")
                        all_outputs[agent_cfg["id"]] = {"raw_text": f"Error: {e}", "agent_id": agent_cfg["id"]}
        else:
            for agent_cfg in phase_agents:
                print(f"  Running {agent_cfg['name']}...")
                try:
                    prompt = _load_prompt(org_id, agent_cfg, state, all_outputs, cycle_date)
                    output = run_agent(org_id, agent_cfg, prompt, cycle_date)
                    all_outputs[agent_cfg["id"]] = output
                    print(f"  ✓ {agent_cfg['name']} complete")
                except Exception as e:
                    print(f"  ✗ {agent_cfg['name']} failed: {e}")
                    all_outputs[agent_cfg["id"]] = {"raw_text": f"Error: {e}", "agent_id": agent_cfg["id"]}

    # Save the report
    reporter_output = all_outputs.get("grant_reporter", {})
    report_text = reporter_output.get("raw_text", "No report generated.")

    # Archive cycle
    from datetime import datetime, timezone
    run_ts = datetime.now(timezone.utc).strftime("%H%M%S")
    cycles_dir = org_cycles_dir(org_id)
    cycle_dir = os.path.join(cycles_dir, cycle_date, f"run-{run_ts}")
    os.makedirs(cycle_dir, exist_ok=True)

    with open(os.path.join(cycle_dir, "report.md"), "w") as f:
        f.write(report_text)

    with open(os.path.join(cycle_dir, "outputs.json"), "w") as f:
        json.dump(all_outputs, f, indent=2, default=str)

    evidence = load_evidence_for_date(org_id, cycle_date)
    with open(os.path.join(cycle_dir, "evidence.json"), "w") as f:
        json.dump(evidence, f, indent=2)

    # Write a "latest" symlink
    latest_dir = os.path.join(cycles_dir, cycle_date, "latest")
    if os.path.islink(latest_dir):
        os.unlink(latest_dir)
    elif os.path.isdir(latest_dir):
        import shutil
        shutil.rmtree(latest_dir)
    os.symlink(f"run-{run_ts}", latest_dir)

    print(f"\n--- Cycle Complete ---")
    print(f"Report: {cycle_dir}/report.md")
    print(f"Evidence items: {len(evidence)}")

    return report_text
