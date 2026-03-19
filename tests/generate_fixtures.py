"""Generate 15 test documents (5 formats x 3 variants) for upload testing."""

import os

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")
os.makedirs(FIXTURES_DIR, exist_ok=True)


def make_pdfs():
    """Generate 3 PDF files."""
    import fitz  # PyMuPDF

    variants = [
        ("simple_form.pdf", "Grant Application Form\n\n"
         "Organisation Name: ___________\n"
         "Project Title: ___________\n"
         "Funding Amount Requested: ___________\n"
         "Project Summary (max 500 words): ___________\n"
         "Expected Outcomes: ___________\n"),
        ("budget template.pdf", "Budget Template\n\n"
         "Personnel Costs: ___________\n"
         "Travel Costs: ___________\n"
         "Equipment: ___________\n"
         "Total Budget: ___________\n"
         "Co-funding Amount: ___________\n"),
        ("Māori Community Grant (2026).pdf", "Māori Community Development Grant\n\n"
         "Applicant Name: ___________\n"
         "Iwi Affiliation: ___________\n"
         "Project Description (max 1000 words): ___________\n"
         "Community Impact: ___________\n"
         "Timeline: ___________\n"
         "Budget Breakdown: ___________\n"),
    ]
    for filename, text in variants:
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), text, fontsize=11)
        path = os.path.join(FIXTURES_DIR, filename)
        doc.save(path)
        doc.close()
        print(f"  Created: {filename}")


def make_docx():
    """Generate 3 DOCX files."""
    from docx import Document

    variants = [
        ("simple_form.docx", [
            ("Organisation Legal Name", ""),
            ("Contact Person", ""),
            ("Project Title", ""),
            ("Project Summary", "Describe your project in no more than 500 words."),
            ("Total Funding Requested", ""),
            ("Project Duration", ""),
        ]),
        ("detailed application.docx", [
            ("Full Name of Applicant Entity", ""),
            ("Registration Number", ""),
            ("Physical Address", ""),
            ("Project Objectives", "List the key objectives of the proposed project."),
            ("Methodology", "Describe your approach in detail."),
            ("Expected Outcomes", "What measurable outcomes do you expect?"),
            ("Risk Assessment", "Identify key risks and mitigation strategies."),
            ("Budget Total (NZD)", ""),
        ]),
        ("Quick Fund Round 3.docx", [
            ("Applicant Name", ""),
            ("Email", ""),
            ("Phone", ""),
            ("Brief Description (max 200 words)", ""),
            ("Amount Requested", ""),
            ("How will you measure success?", ""),
        ]),
    ]
    for filename, fields in variants:
        doc = Document()
        doc.add_heading(filename.replace(".docx", ""), level=0)
        table = doc.add_table(rows=len(fields), cols=2)
        table.style = "Table Grid"
        for i, (label, hint) in enumerate(fields):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = hint if hint else "[TO BE COMPLETED]"
        path = os.path.join(FIXTURES_DIR, filename)
        doc.save(path)
        print(f"  Created: {filename}")


def make_doc():
    """Generate 3 .doc files (actually RTF that Word can open, or plain text with .doc extension).
    Since we can't create real OLE .doc without MS Word, create plain text .doc files
    which textutil on macOS can still read."""
    variants = [
        ("legacy_form.doc",
         "Grant Application Form - Legacy Format\n\n"
         "Organisation Name: ___________\n"
         "ABN/NZBN: ___________\n"
         "Project Title: ___________\n"
         "Funding Amount: ___________\n"
         "Project Description: ___________\n"),
        ("old budget form.doc",
         "Budget Submission Form\n\n"
         "Item 1: Personnel - Amount: ___________\n"
         "Item 2: Travel - Amount: ___________\n"
         "Item 3: Equipment - Amount: ___________\n"
         "Item 4: Other - Amount: ___________\n"
         "Total: ___________\n"),
        ("Community Impact (v1).doc",
         "Community Impact Assessment Form\n\n"
         "Organisation: ___________\n"
         "Target Community: ___________\n"
         "Number of Beneficiaries: ___________\n"
         "Impact Description: ___________\n"
         "Sustainability Plan: ___________\n"),
    ]
    for filename, text in variants:
        # Write as RTF so textutil can convert it
        rtf = r"{\rtf1\ansi " + text.replace("\n", r" \par ") + "}"
        path = os.path.join(FIXTURES_DIR, filename)
        with open(path, "w") as f:
            f.write(rtf)
        print(f"  Created: {filename}")


def make_xlsx():
    """Generate 3 XLSX files."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    variants = [
        ("budget_template.xlsx", "Budget", [
            ["Category", "Description", "Amount (NZD)"],
            ["Personnel", "", ""],
            ["Travel", "", ""],
            ["Equipment", "", ""],
            ["Consumables", "", ""],
            ["Other", "", ""],
            ["TOTAL", "", "=SUM(C2:C7)"],
        ]),
        ("monitoring framework.xlsx", "M&E Framework", [
            ["Indicator", "Baseline", "Target", "Data Source"],
            ["Number of participants trained", "", "", ""],
            ["Community satisfaction score", "", "", ""],
            ["Revenue generated (NZD)", "", "", ""],
            ["Jobs created", "", "", ""],
        ]),
        ("Grant Financials (2026).xlsx", "Financials", [
            ["Line Item", "Q1", "Q2", "Q3", "Q4", "Total"],
            ["Staff costs", "", "", "", "", ""],
            ["Operational costs", "", "", "", "", ""],
            ["Capital expenditure", "", "", "", "", ""],
            ["Contingency", "", "", "", "", ""],
            ["Grand Total", "", "", "", "", ""],
        ]),
    ]
    for filename, sheet_name, rows in variants:
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name
        for i, row in enumerate(rows):
            ws.append(row)
            if i == 0:
                for cell in ws[1]:
                    cell.font = Font(bold=True)
        path = os.path.join(FIXTURES_DIR, filename)
        wb.save(path)
        print(f"  Created: {filename}")


def make_xls():
    """Generate 3 XLS files."""
    import xlrd
    # xlrd is read-only; use openpyxl to create xlsx then convert won't work.
    # Instead, create minimal XLS files using the struct directly.
    # Simplest: write as CSV with .xls extension (Excel opens these).
    # For proper testing, let's use xlwt if available, otherwise CSV fallback.
    try:
        import xlwt
        _make_xls_xlwt()
    except ImportError:
        _make_xls_csv_fallback()


def _make_xls_xlwt():
    import xlwt

    variants = [
        ("legacy_budget.xls", "Budget", [
            ["Category", "Amount"],
            ["Staff", ""],
            ["Travel", ""],
            ["Equipment", ""],
            ["Total", ""],
        ]),
        ("old financials.xls", "Sheet1", [
            ["Item", "Cost (NZD)", "Notes"],
            ["Venue hire", "", ""],
            ["Catering", "", ""],
            ["Materials", "", ""],
            ["Facilitator fees", "", ""],
        ]),
        ("Quarterly Report (legacy).xls", "Q1 Report", [
            ["Metric", "Target", "Actual"],
            ["Participants", "100", ""],
            ["Workshops delivered", "4", ""],
            ["Satisfaction rating", "4.0", ""],
        ]),
    ]
    for filename, sheet_name, rows in variants:
        wb = xlwt.Workbook()
        ws = wb.add_sheet(sheet_name)
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                ws.write(r, c, val)
        path = os.path.join(FIXTURES_DIR, filename)
        wb.save(path)
        print(f"  Created: {filename}")


def _make_xls_csv_fallback():
    """Fallback: create CSV files with .xls extension."""
    variants = [
        ("legacy_budget.xls", [
            "Category,Amount",
            "Staff,",
            "Travel,",
            "Equipment,",
            "Total,",
        ]),
        ("old financials.xls", [
            "Item,Cost (NZD),Notes",
            "Venue hire,,",
            "Catering,,",
            "Materials,,",
            "Facilitator fees,,",
        ]),
        ("Quarterly Report (legacy).xls", [
            "Metric,Target,Actual",
            "Participants,100,",
            "Workshops delivered,4,",
            "Satisfaction rating,4.0,",
        ]),
    ]
    for filename, rows in variants:
        path = os.path.join(FIXTURES_DIR, filename)
        with open(path, "w") as f:
            f.write("\n".join(rows))
        print(f"  Created: {filename} (CSV fallback)")


if __name__ == "__main__":
    print("Generating PDF files...")
    make_pdfs()
    print("Generating DOCX files...")
    make_docx()
    print("Generating DOC files...")
    make_doc()
    print("Generating XLSX files...")
    make_xlsx()
    print("Generating XLS files...")
    make_xls()
    print(f"\nDone! {len(os.listdir(FIXTURES_DIR))} files in {FIXTURES_DIR}")
