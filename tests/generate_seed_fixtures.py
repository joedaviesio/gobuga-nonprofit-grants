"""Generate sample org documents for seeding upload tests (5 doc types)."""

import os

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "seed_fixtures")
os.makedirs(FIXTURES_DIR, exist_ok=True)


def make_annual_report():
    """PDF annual report."""
    import fitz

    text = (
        "Annual Report 2025 — Kiwi Community Trust\n\n"
        "Chair's Report\n"
        "This year marked significant growth for Kiwi Community Trust. We expanded "
        "our digital inclusion programme to 12 new communities across Northland and "
        "the Bay of Plenty, reaching over 3,200 individuals.\n\n"
        "Financial Highlights\n"
        "Total Revenue: $1,245,000\n"
        "Government Grants: $680,000\n"
        "Philanthropic Funding: $320,000\n"
        "Earned Income: $145,000\n"
        "Donations: $100,000\n"
        "Total Expenditure: $1,180,000\n"
        "Net Surplus: $65,000\n\n"
        "Key Achievements\n"
        "- 3,200 participants in digital literacy workshops\n"
        "- 450 seniors connected online for the first time\n"
        "- 89% participant satisfaction rate\n"
        "- Partnered with 8 local councils\n"
        "- Launched te reo Māori digital resources\n\n"
        "Looking Ahead\n"
        "In 2026 we plan to expand into Waikato and Taranaki regions, "
        "with a target of 5,000 participants and a new youth coding programme."
    )
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), text, fontsize=10)
    path = os.path.join(FIXTURES_DIR, "annual-report-2025.pdf")
    doc.save(path)
    doc.close()
    print(f"  Created: annual-report-2025.pdf")


def make_mission_statement():
    """DOCX mission statement."""
    from docx import Document

    doc = Document()
    doc.add_heading("Kiwi Community Trust — Mission Statement", level=0)
    doc.add_paragraph(
        "Our mission is to bridge the digital divide in Aotearoa New Zealand by "
        "providing accessible, culturally responsive technology education and "
        "connectivity solutions to underserved communities."
    )
    doc.add_heading("Vision", level=1)
    doc.add_paragraph(
        "A digitally inclusive Aotearoa where every person, regardless of age, "
        "location, or background, can participate fully in the digital world."
    )
    doc.add_heading("Values", level=1)
    for v in [
        "Manaakitanga — We care for our communities with kindness and respect",
        "Kotahitanga — We work together as one",
        "Kaitiakitanga — We are stewards of our resources and knowledge",
        "Equity — We prioritise those with the greatest need",
        "Innovation — We embrace new approaches to old challenges",
    ]:
        doc.add_paragraph(v, style="List Bullet")
    doc.add_heading("Strategic Goals 2025-2028", level=1)
    for g in [
        "Reach 10,000 participants annually by 2028",
        "Establish presence in all 16 regions of NZ",
        "Develop bilingual (English/te reo Māori) digital curriculum",
        "Achieve Tier 2 charitable status and $2M annual revenue",
        "Launch social enterprise arm for sustainable income",
    ]:
        doc.add_paragraph(g, style="List Number")

    path = os.path.join(FIXTURES_DIR, "mission-statement.docx")
    doc.save(path)
    print(f"  Created: mission-statement.docx")


def make_org_review():
    """Text-based org review document."""
    text = (
        "Organisational Review — Kiwi Community Trust\n"
        "Conducted by: Independent Governance Advisory, March 2025\n\n"
        "Executive Summary\n"
        "Kiwi Community Trust is a well-managed charitable organisation with "
        "strong community relationships and growing impact. This review "
        "identifies areas for improvement in financial sustainability, "
        "governance practices, and operational efficiency.\n\n"
        "Strengths\n"
        "- Deeply trusted by target communities across Northland and BOP\n"
        "- Strong volunteer base of 120+ regular volunteers\n"
        "- Effective partnerships with local councils and iwi\n"
        "- High programme completion rates (87%)\n"
        "- Experienced leadership team with sector expertise\n\n"
        "Areas for Improvement\n"
        "- Revenue concentration: 55% from government grants\n"
        "- Board composition lacks financial expertise\n"
        "- M&E framework needs strengthening\n"
        "- IT systems require modernisation\n"
        "- Staff retention below sector average\n\n"
        "Recommendations\n"
        "1. Diversify funding — target 40% government, 30% philanthropy, "
        "20% earned income, 10% donations by 2027\n"
        "2. Recruit board member with CFO/accounting background\n"
        "3. Implement outcomes measurement framework (Theory of Change)\n"
        "4. Invest in cloud-based CRM and programme management tools\n"
        "5. Develop staff professional development programme\n"
    )
    path = os.path.join(FIXTURES_DIR, "organisational-review-2025.txt")
    with open(path, "w") as f:
        f.write(text)
    print(f"  Created: organisational-review-2025.txt")


def make_previous_application():
    """DOCX previous grant application."""
    from docx import Document

    doc = Document()
    doc.add_heading("Lottery Grants Board Application — Digital Inclusion NZ", level=0)
    doc.add_paragraph("Applicant: Kiwi Community Trust")
    doc.add_paragraph("Date Submitted: 15 September 2024")
    doc.add_paragraph("Amount Requested: $185,000")
    doc.add_paragraph("Status: Approved ($150,000)")

    doc.add_heading("Project Summary", level=1)
    doc.add_paragraph(
        "Digital Inclusion NZ will deliver a 12-month programme of digital "
        "literacy workshops, device lending, and internet connectivity support "
        "across 8 communities in Northland. The programme targets seniors (65+), "
        "new migrants, and rural youth who face significant digital exclusion."
    )

    doc.add_heading("Objectives", level=1)
    for obj in [
        "Train 800 participants in basic digital skills",
        "Distribute 200 refurbished laptops to families in need",
        "Establish 6 community WiFi hotspots in underserved areas",
        "Develop 3 te reo Māori digital learning modules",
    ]:
        doc.add_paragraph(obj, style="List Number")

    doc.add_heading("Budget Summary", level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    data = [
        ("Item", "Amount (NZD)"),
        ("Programme Coordinator (1 FTE)", "$75,000"),
        ("Workshop Facilitators (3 PT)", "$54,000"),
        ("Equipment & Devices", "$30,000"),
        ("Travel & Venue Hire", "$16,000"),
        ("Total", "$175,000"),
    ]
    for i, (label, amount) in enumerate(data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = amount

    doc.add_heading("Outcomes Achieved (Final Report)", level=1)
    doc.add_paragraph(
        "The programme exceeded most targets: 920 participants trained (115%), "
        "230 devices distributed, 7 WiFi hotspots established. Te reo modules "
        "were completed and are now freely available online."
    )

    path = os.path.join(FIXTURES_DIR, "previous-application-lottery-2024.docx")
    doc.save(path)
    print(f"  Created: previous-application-lottery-2024.docx")


def make_financial_statement():
    """XLSX financial statement."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, numbers

    wb = Workbook()
    ws = wb.active
    ws.title = "Income & Expenditure"

    ws.append(["Kiwi Community Trust — Financial Statement FY2025"])
    ws.merge_cells("A1:C1")
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([])

    ws.append(["INCOME", "", "Amount (NZD)"])
    ws["A3"].font = Font(bold=True)
    ws["C3"].font = Font(bold=True)
    income = [
        ("Government Grants — MSD", 320000),
        ("Government Grants — DIA", 180000),
        ("Government Grants — MBIE", 180000),
        ("Foundation North", 120000),
        ("Todd Foundation", 85000),
        ("Tindall Foundation", 65000),
        ("Corporate Sponsorships", 50000),
        ("Fee for Service", 95000),
        ("Individual Donations", 78000),
        ("Interest Income", 12000),
    ]
    for label, amount in income:
        ws.append([label, "", amount])
    total_income_row = ws.max_row + 1
    ws.append(["TOTAL INCOME", "", f"=SUM(C4:C{total_income_row-1})"])
    ws[f"A{total_income_row}"].font = Font(bold=True)
    ws.append([])

    ws.append(["EXPENDITURE", "", "Amount (NZD)"])
    ws[f"A{ws.max_row}"].font = Font(bold=True)
    expenditure = [
        ("Salaries & Wages", 520000),
        ("Programme Delivery", 185000),
        ("Travel & Transport", 48000),
        ("Equipment & Technology", 95000),
        ("Office & Administration", 62000),
        ("Professional Development", 28000),
        ("Insurance", 18000),
        ("Audit & Legal", 22000),
        ("Communications & Marketing", 35000),
        ("Depreciation", 42000),
    ]
    exp_start = ws.max_row + 1
    for label, amount in expenditure:
        ws.append([label, "", amount])
    total_exp_row = ws.max_row + 1
    ws.append(["TOTAL EXPENDITURE", "", f"=SUM(C{exp_start}:C{total_exp_row-1})"])
    ws[f"A{total_exp_row}"].font = Font(bold=True)
    ws.append([])
    ws.append(["NET SURPLUS/(DEFICIT)", "", f"=C{total_income_row}-C{total_exp_row}"])
    ws[f"A{ws.max_row}"].font = Font(bold=True, color="006600")

    ws.column_dimensions["A"].width = 35
    ws.column_dimensions["C"].width = 18

    path = os.path.join(FIXTURES_DIR, "financial-statement-fy2025.xlsx")
    wb.save(path)
    print(f"  Created: financial-statement-fy2025.xlsx")


if __name__ == "__main__":
    print("Generating seed fixture documents...\n")
    make_annual_report()
    make_mission_statement()
    make_org_review()
    make_previous_application()
    make_financial_statement()
    print(f"\nDone! {len(os.listdir(FIXTURES_DIR))} files in {FIXTURES_DIR}")
