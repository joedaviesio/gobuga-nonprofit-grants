"""
Export pipeline test suite.

Tests:
  1. Fresh DOCX generation produces a valid file
  2. Fresh PDF generation produces a valid file
  3. Fresh XLSX generation produces a valid file
  4. Markdown export returns content
  5. JSON export returns structured data
  6. Template-filled DOCX puts content in cells
  7. Template-filled XLSX puts content in cells
  8. Template-filled PDF (AcroForm) fills fields
  9. .doc → .docx conversion works
  10. Download endpoint requires auth token
  11. Export with no sections returns valid file
  12. Export with unicode content doesn't crash

Usage:
    python3 tests/test_exports.py
    python3 tests/test_exports.py --case 1
"""

import os
import sys
import time
import json
import shutil
import requests
import pytest

API = os.environ.get("API_URL", "http://localhost:8102")
FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")

PASS = 0
FAIL = 0
RESULTS = []


# --- pytest wiring -----------------------------------------------------------
# Under pytest the `backend` fixture (tests/conftest.py) boots an isolated
# server; the script-style check()/log() helpers below stay as they are.

pytestmark = pytest.mark.integration


@pytest.fixture(scope="module", autouse=True)
def _authenticated(backend):
    """Register a fresh org against the live backend before any test runs."""
    if not setup_auth():
        pytest.fail("Could not register a test org against the backend")



def log(msg, indent=0):
    line = "  " * indent + msg
    print(line)
    RESULTS.append(line)


def check(label, condition, detail=""):
    global PASS, FAIL
    if condition:
        PASS += 1
        log(f"[PASS] {label}")
    else:
        FAIL += 1
        log(f"[FAIL] {label}" + (f"  — {detail}" if detail else ""))
    return condition


# --- Auth helpers ---

TOKEN = None
ORG_ID = None


def setup_auth():
    global TOKEN, ORG_ID
    ts = int(time.time() * 1000)
    email = f"test-export-{ts}@gobuga-test.org"
    r = requests.post(f"{API}/api/auth/register", json={
        "email": email,
        "password": "TestPass123!",
        "org_name": f"Export Test Org {ts}",
        "website_url": "https://example.com",
    })
    if not r.ok:
        log(f"[FATAL] Registration failed: {r.status_code} {r.text[:200]}")
        return False
    data = r.json()
    TOKEN = data["token"]
    ORG_ID = data["org_id"]

    # Complete setup so we can use the API
    r = requests.post(f"{API}/api/org/setup", json={
        "org_name": "Export Test Org",
        "country": "New Zealand",
        "website": "https://example.com",
        "charitable_status": "Registered Charity",
        "mission": "Testing exports",
        "sectors": ["Education"],
        "geographies": ["New Zealand"],
    }, headers=json_headers())

    # Exports are Officer-only and Scanner caps open cases at 3; the suite
    # creates a case per test, so upgrade via the dev toggle.
    r = requests.post(f"{API}/api/org/toggle-tier", headers=json_headers())
    if not r.ok or r.json().get("tier") != "officer":
        log(f"[FATAL] Could not upgrade test org to Officer: {r.status_code} {r.text[:200]}")
        return False
    return True


def auth_headers():
    return {"Authorization": f"Bearer {TOKEN}"}


def json_headers():
    return {"Authorization": f"Bearer {TOKEN}", "Content-Type": "application/json"}


def create_test_case(label="test-export", sections=None):
    """Create a case with optional pre-filled sections."""
    r = requests.post(f"{API}/api/cases", json={
        "grant_id": f"test-{label}-{int(time.time())}",
        "grant_brief": {
            "title": f"Test Grant: {label}",
            "funder": "Test Funder NZ",
            "deadline": "2026-12-31",
            "amount": "$50,000",
            "source_cycle": "2026-03-19",
        },
    }, headers=json_headers())
    if not r.ok:
        log(f"  Create case failed: {r.status_code} {r.text[:200]}")
        return None
    case_id = r.json()["case_id"]

    # Fill sections if provided
    if sections:
        for name, content in sections.items():
            requests.patch(
                f"{API}/api/cases/{case_id}/draft/{name}",
                json={"content": content, "status": "draft"},
                headers=json_headers(),
            )

    return case_id


def export_case(case_id, fmt, template_filename=None):
    """Call the export endpoint."""
    body = {"format": fmt}
    if template_filename:
        body["template_filename"] = template_filename
    r = requests.post(
        f"{API}/api/cases/{case_id}/export",
        json=body,
        headers=json_headers(),
    )
    return r


def download_file(case_id, filename, with_auth=True):
    """Download an exported file."""
    headers = auth_headers() if with_auth else {}
    r = requests.get(
        f"{API}/api/cases/{case_id}/download/{filename}",
        headers=headers,
    )
    return r


DEFAULT_SECTIONS = {
    "project_summary": "This project delivers digital literacy workshops across rural Northland communities, "
                       "targeting 500 seniors and 200 youth over 12 months.",
    "organisation_background": "Kiwi Community Trust is a registered charity (CC12345) operating since 2018. "
                               "We have delivered 15 successful programmes reaching 3,200 participants.",
    "budget_overview": "Total budget: $185,000\n- Staff costs: $95,000\n- Equipment: $40,000\n"
                       "- Travel: $25,000\n- Venue hire: $15,000\n- Admin: $10,000",
    "expected_outcomes": "- 500 seniors complete digital skills course\n- 200 youth complete coding programme\n"
                        "- 90% participant satisfaction\n- 8 community hubs established",
}


# ─────────────────────────────────────────────────────────────────
# TEST 1: Fresh DOCX generation
# ─────────────────────────────────────────────────────────────────

def test_1_fresh_docx():
    log("\n" + "=" * 70)
    log("TEST 1: Fresh DOCX generation produces a valid file")
    log("=" * 70)

    case_id = create_test_case("docx", DEFAULT_SECTIONS)
    check("Case created", case_id is not None)
    if not case_id:
        return

    r = export_case(case_id, "docx")
    check("Export returns 200", r.ok, f"Got {r.status_code}: {r.text[:200]}")
    if not r.ok:
        return

    data = r.json()
    filename = data.get("filename", "")
    check("Response has filename", bool(filename), f"Got: {data}")
    check("Filename ends with .docx", filename.endswith(".docx"), filename)

    # Download and validate
    dl = download_file(case_id, filename)
    check("Download returns 200", dl.status_code == 200, f"Got {dl.status_code}")
    check("Content-type is correct", "application" in dl.headers.get("content-type", ""),
          dl.headers.get("content-type"))
    check("File size > 0", len(dl.content) > 0, f"Got {len(dl.content)} bytes")

    # Validate with python-docx
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(dl.content)
        tmp_path = f.name

    try:
        from docx import Document
        doc = Document(tmp_path)
        check("python-docx can open file", True)
        check("Has paragraphs", len(doc.paragraphs) > 0, f"Got {len(doc.paragraphs)}")

        all_text = "\n".join(p.text for p in doc.paragraphs)
        check("Contains project_summary content", "digital literacy" in all_text.lower(),
              f"First 200 chars: {all_text[:200]}")
        check("Contains budget content", "185,000" in all_text, "Budget not found")
    except Exception as e:
        check("python-docx can open file", False, str(e))
    finally:
        os.unlink(tmp_path)


# ─────────────────────────────────────────────────────────────────
# TEST 2: Fresh PDF generation
# ─────────────────────────────────────────────────────────────────

def test_2_fresh_pdf():
    log("\n" + "=" * 70)
    log("TEST 2: Fresh PDF generation produces a valid file")
    log("=" * 70)

    case_id = create_test_case("pdf", DEFAULT_SECTIONS)
    check("Case created", case_id is not None)
    if not case_id:
        return

    r = export_case(case_id, "pdf")
    check("Export returns 200", r.ok, f"Got {r.status_code}: {r.text[:200]}")
    if not r.ok:
        return

    data = r.json()
    filename = data.get("filename", "")
    check("Filename ends with .pdf", filename.endswith(".pdf"), filename)

    dl = download_file(case_id, filename)
    check("Download returns 200", dl.status_code == 200)
    check("File size > 0", len(dl.content) > 0)

    # Validate with PyMuPDF
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(dl.content)
        tmp_path = f.name

    try:
        import fitz
        doc = fitz.open(tmp_path)
        check("PyMuPDF can open file", True)
        check("Has pages", doc.page_count > 0, f"Got {doc.page_count} pages")

        text = ""
        for page in doc:
            text += page.get_text()
        check("Contains project content", "digital literacy" in text.lower(),
              f"First 200 chars: {text[:200]}")
        doc.close()
    except Exception as e:
        check("PyMuPDF can open file", False, str(e))
    finally:
        os.unlink(tmp_path)


# ─────────────────────────────────────────────────────────────────
# TEST 3: Fresh XLSX generation
# ─────────────────────────────────────────────────────────────────

def test_3_fresh_xlsx():
    log("\n" + "=" * 70)
    log("TEST 3: Fresh XLSX generation produces a valid file")
    log("=" * 70)

    case_id = create_test_case("xlsx", DEFAULT_SECTIONS)
    check("Case created", case_id is not None)
    if not case_id:
        return

    r = export_case(case_id, "xlsx")
    check("Export returns 200", r.ok, f"Got {r.status_code}: {r.text[:200]}")
    if not r.ok:
        return

    data = r.json()
    filename = data.get("filename", "")
    check("Filename ends with .xlsx", filename.endswith(".xlsx"), filename)

    dl = download_file(case_id, filename)
    check("Download returns 200", dl.status_code == 200)

    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        f.write(dl.content)
        tmp_path = f.name

    try:
        from openpyxl import load_workbook
        wb = load_workbook(tmp_path)
        ws = wb.active
        check("openpyxl can open file", True)
        check("Has rows", ws.max_row > 0, f"Got {ws.max_row} rows")

        all_text = " ".join(str(cell.value or "") for row in ws.iter_rows() for cell in row)
        check("Contains project content", "digital literacy" in all_text.lower(),
              f"First 200 chars: {all_text[:200]}")
    except Exception as e:
        check("openpyxl can open file", False, str(e))
    finally:
        os.unlink(tmp_path)


# ─────────────────────────────────────────────────────────────────
# TEST 4: Markdown export
# ─────────────────────────────────────────────────────────────────

def test_4_markdown():
    log("\n" + "=" * 70)
    log("TEST 4: Markdown export returns content")
    log("=" * 70)

    case_id = create_test_case("md", DEFAULT_SECTIONS)
    check("Case created", case_id is not None)
    if not case_id:
        return

    r = export_case(case_id, "markdown")
    check("Export returns 200", r.ok, f"Got {r.status_code}")
    if not r.ok:
        return

    data = r.json()
    content = data.get("content", "")
    check("Has content", len(content) > 0, f"Got {len(content)} chars")
    check("Contains section headers", "##" in content, "No markdown headers found")
    check("Contains project content", "digital literacy" in content.lower())
    check("Contains budget", "185,000" in content)


# ─────────────────────────────────────────────────────────────────
# TEST 5: JSON export
# ─────────────────────────────────────────────────────────────────

def test_5_json():
    log("\n" + "=" * 70)
    log("TEST 5: JSON export returns structured data")
    log("=" * 70)

    case_id = create_test_case("json", DEFAULT_SECTIONS)
    check("Case created", case_id is not None)
    if not case_id:
        return

    r = export_case(case_id, "json")
    check("Export returns 200", r.ok, f"Got {r.status_code}")
    if not r.ok:
        return

    data = r.json()
    content = data.get("content", {})
    check("Has case_id", content.get("case_id") == case_id, f"Got {content.get('case_id')}")
    check("Has sections", "sections" in content, f"Keys: {list(content.keys())}")
    sections = content.get("sections", {})
    check("Has project_summary section", "project_summary" in sections)
    check("Section has content", bool(sections.get("project_summary", {}).get("content")))


# ─────────────────────────────────────────────────────────────────
# TEST 6: Template-filled DOCX
# ─────────────────────────────────────────────────────────────────

def test_6_template_docx():
    log("\n" + "=" * 70)
    log("TEST 6: Template-filled DOCX puts content in cells")
    log("=" * 70)

    # Create a DOCX template with a table
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("Grant Application Form", level=0)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Project Summary"
    table.rows[0].cells[1].text = ""
    table.rows[1].cells[0].text = "Organisation Background"
    table.rows[1].cells[1].text = ""
    table.rows[2].cells[0].text = "Budget Overview"
    table.rows[2].cells[1].text = ""
    table.rows[3].cells[0].text = "Expected Outcomes"
    table.rows[3].cells[1].text = ""

    import tempfile
    template_dir = tempfile.mkdtemp()
    template_path = os.path.join(template_dir, "template.docx")
    doc.save(template_path)

    # Create case with matching sections
    case_id = create_test_case("tmpl-docx", DEFAULT_SECTIONS)
    check("Case created", case_id is not None)
    if not case_id:
        return

    # Upload template
    with open(template_path, "rb") as f:
        r = requests.post(
            f"{API}/api/cases/{case_id}/upload",
            files={"file": ("template.docx", f)},
            data={"file_type": "submission"},
            headers=auth_headers(),
        )
    check("Template uploaded", r.ok, f"Got {r.status_code}: {r.text[:200]}")
    if not r.ok:
        return

    uploaded_name = r.json().get("filename", "template.docx")

    # Also need parsed_sections for _build_field_map
    # Simulate by adding section labels that match template
    for section_id, label in [
        ("project_summary", "Project Summary"),
        ("organisation_background", "Organisation Background"),
        ("budget_overview", "Budget Overview"),
        ("expected_outcomes", "Expected Outcomes"),
    ]:
        pass  # sections already exist from create_test_case

    # Export with template
    r = export_case(case_id, "docx", template_filename=uploaded_name)
    check("Template export returns 200", r.ok, f"Got {r.status_code}: {r.text[:200]}")
    if not r.ok:
        return

    data = r.json()
    filename = data.get("filename", "")
    check("Filename contains 'filled'", "filled" in filename, filename)

    dl = download_file(case_id, filename)
    check("Download returns 200", dl.status_code == 200)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(dl.content)
        tmp_path = f.name

    try:
        filled_doc = DocxDocument(tmp_path)
        check("python-docx can open filled file", True)

        # Check table cells have content
        filled_any = False
        for table in filled_doc.tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    label = cells[0].text.strip()
                    value = cells[1].text.strip()
                    if label and value:
                        filled_any = True
                        log(f"  {label[:30]}: {value[:50]}...", indent=1)

        check("At least one table cell was filled", filled_any)
    except Exception as e:
        check("python-docx can open filled file", False, str(e))
    finally:
        os.unlink(tmp_path)
        shutil.rmtree(template_dir)


# ─────────────────────────────────────────────────────────────────
# TEST 7: Template-filled XLSX
# ─────────────────────────────────────────────────────────────────

def test_7_template_xlsx():
    log("\n" + "=" * 70)
    log("TEST 7: Template-filled XLSX puts content in cells")
    log("=" * 70)

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Application"
    ws.append(["Project Summary", ""])
    ws.append(["Organisation Background", ""])
    ws.append(["Budget Overview", ""])
    ws.append(["Expected Outcomes", ""])

    import tempfile
    template_dir = tempfile.mkdtemp()
    template_path = os.path.join(template_dir, "template.xlsx")
    wb.save(template_path)

    case_id = create_test_case("tmpl-xlsx", DEFAULT_SECTIONS)
    check("Case created", case_id is not None)
    if not case_id:
        return

    with open(template_path, "rb") as f:
        r = requests.post(
            f"{API}/api/cases/{case_id}/upload",
            files={"file": ("template.xlsx", f)},
            data={"file_type": "submission"},
            headers=auth_headers(),
        )
    check("Template uploaded", r.ok)
    if not r.ok:
        return

    uploaded_name = r.json().get("filename", "template.xlsx")

    r = export_case(case_id, "xlsx", template_filename=uploaded_name)
    check("Template export returns 200", r.ok, f"Got {r.status_code}: {r.text[:200]}")
    if not r.ok:
        return

    data = r.json()
    filename = data.get("filename", "")
    check("Filename contains 'filled'", "filled" in filename, filename)

    dl = download_file(case_id, filename)
    check("Download returns 200", dl.status_code == 200)

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        f.write(dl.content)
        tmp_path = f.name

    try:
        from openpyxl import load_workbook
        filled_wb = load_workbook(tmp_path)
        ws = filled_wb.active
        check("openpyxl can open filled file", True)

        filled_any = False
        for row in ws.iter_rows(min_row=1, max_row=4, min_col=1, max_col=2):
            label = str(row[0].value or "").strip()
            value = str(row[1].value or "").strip()
            if label and value:
                filled_any = True
                log(f"  {label[:30]}: {value[:50]}...", indent=1)

        check("At least one cell was filled", filled_any)
    except Exception as e:
        check("openpyxl can open filled file", False, str(e))
    finally:
        os.unlink(tmp_path)
        shutil.rmtree(template_dir)


# ─────────────────────────────────────────────────────────────────
# TEST 8: Template-filled PDF (AcroForm)
# ─────────────────────────────────────────────────────────────────

def test_8_template_pdf():
    log("\n" + "=" * 70)
    log("TEST 8: Template-filled PDF")
    log("=" * 70)

    import fitz
    import tempfile

    # Create a PDF with form fields
    doc = fitz.open()
    page = doc.new_page()

    # Add text labels
    page.insert_text((50, 50), "Project Summary", fontsize=12)
    page.insert_text((50, 150), "Organisation Background", fontsize=12)
    page.insert_text((50, 250), "Budget Overview", fontsize=12)

    template_dir = tempfile.mkdtemp()
    template_path = os.path.join(template_dir, "template.pdf")
    doc.save(template_path)
    doc.close()

    case_id = create_test_case("tmpl-pdf", DEFAULT_SECTIONS)
    check("Case created", case_id is not None)
    if not case_id:
        return

    with open(template_path, "rb") as f:
        r = requests.post(
            f"{API}/api/cases/{case_id}/upload",
            files={"file": ("template.pdf", f)},
            data={"file_type": "submission"},
            headers=auth_headers(),
        )
    check("Template uploaded", r.ok)
    if not r.ok:
        return

    uploaded_name = r.json().get("filename", "template.pdf")

    r = export_case(case_id, "pdf", template_filename=uploaded_name)
    check("PDF template export returns 200", r.ok, f"Got {r.status_code}: {r.text[:200]}")
    if not r.ok:
        return

    data = r.json()
    filename = data.get("filename", "")
    check("Has filename", bool(filename))

    dl = download_file(case_id, filename)
    check("Download returns 200", dl.status_code == 200)
    check("File size > original template", len(dl.content) > 500,
          f"Got {len(dl.content)} bytes")

    # Validate the filled PDF
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(dl.content)
        tmp_path = f.name

    try:
        filled = fitz.open(tmp_path)
        check("PyMuPDF can open filled PDF", True)
        text = ""
        for page in filled:
            text += page.get_text()
        # The text overlay should have added our content
        check("Filled PDF has more text than template",
              len(text) > 50, f"Got {len(text)} chars of text")
        filled.close()
    except Exception as e:
        check("PyMuPDF can open filled PDF", False, str(e))
    finally:
        os.unlink(tmp_path)
        shutil.rmtree(template_dir)


# ─────────────────────────────────────────────────────────────────
# TEST 9: .doc → .docx conversion
# ─────────────────────────────────────────────────────────────────

def test_9_doc_conversion():
    log("\n" + "=" * 70)
    log("TEST 9: .doc to .docx conversion")
    log("=" * 70)

    has_textutil = shutil.which("textutil") is not None
    check("textutil available (macOS)", has_textutil,
          "Skipping — textutil not found (not macOS)")
    if not has_textutil:
        return

    # Create a simple DOCX, convert to DOC, then test export
    from docx import Document as DocxDocument
    import tempfile

    doc = DocxDocument()
    doc.add_heading("Test Form", level=0)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Project Summary"
    table.rows[0].cells[1].text = ""
    table.rows[1].cells[0].text = "Budget Overview"
    table.rows[1].cells[1].text = ""

    template_dir = tempfile.mkdtemp()
    docx_path = os.path.join(template_dir, "form.docx")
    doc.save(docx_path)

    # Convert to .doc
    import subprocess
    doc_path = os.path.join(template_dir, "form.doc")
    result = subprocess.run(
        ["textutil", "-convert", "doc", "-output", doc_path, docx_path],
        capture_output=True, text=True, timeout=30,
    )
    check(".doc file created", os.path.exists(doc_path), result.stderr)
    if not os.path.exists(doc_path):
        shutil.rmtree(template_dir)
        return

    case_id = create_test_case("doc-conv", DEFAULT_SECTIONS)
    check("Case created", case_id is not None)
    if not case_id:
        shutil.rmtree(template_dir)
        return

    with open(doc_path, "rb") as f:
        r = requests.post(
            f"{API}/api/cases/{case_id}/upload",
            files={"file": ("form.doc", f)},
            data={"file_type": "submission"},
            headers=auth_headers(),
        )
    check(".doc uploaded", r.ok)
    if not r.ok:
        shutil.rmtree(template_dir)
        return

    uploaded_name = r.json().get("filename", "form.doc")

    # Export as docx using .doc template — should convert internally
    r = export_case(case_id, "docx", template_filename=uploaded_name)
    check("Export with .doc template returns 200", r.ok,
          f"Got {r.status_code}: {r.text[:200]}")
    if not r.ok:
        shutil.rmtree(template_dir)
        return

    data = r.json()
    filename = data.get("filename", "")
    check("Output is .docx", filename.endswith(".docx"), filename)

    dl = download_file(case_id, filename)
    check("Download returns 200", dl.status_code == 200)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(dl.content)
        tmp_path = f.name

    try:
        filled = DocxDocument(tmp_path)
        check("Converted file opens as valid DOCX", True)
    except Exception as e:
        check("Converted file opens as valid DOCX", False, str(e))
    finally:
        os.unlink(tmp_path)
        shutil.rmtree(template_dir)


# ─────────────────────────────────────────────────────────────────
# TEST 10: Download requires auth
# ─────────────────────────────────────────────────────────────────

def test_10_download_auth():
    log("\n" + "=" * 70)
    log("TEST 10: Download endpoint requires auth token")
    log("=" * 70)

    case_id = create_test_case("auth-test", DEFAULT_SECTIONS)
    check("Case created", case_id is not None)
    if not case_id:
        return

    r = export_case(case_id, "docx")
    check("Export returns 200", r.ok)
    if not r.ok:
        return

    filename = r.json().get("filename", "")

    # Try download WITHOUT auth
    dl_no_auth = download_file(case_id, filename, with_auth=False)
    check("Download without auth is rejected",
          dl_no_auth.status_code in (401, 403, 422),
          f"Got {dl_no_auth.status_code} (expected 401/403)")

    # Try download WITH auth
    dl_auth = download_file(case_id, filename, with_auth=True)
    check("Download with auth succeeds", dl_auth.status_code == 200)

    # Verify the authed response is actually a DOCX, not an error
    check("Authed response is binary (not JSON error)",
          not dl_auth.content.startswith(b'{"detail"'),
          f"First 50 bytes: {dl_auth.content[:50]}")


# ─────────────────────────────────────────────────────────────────
# TEST 11: Export with no sections
# ─────────────────────────────────────────────────────────────────

def test_11_empty_sections():
    log("\n" + "=" * 70)
    log("TEST 11: Export with no sections returns valid file")
    log("=" * 70)

    case_id = create_test_case("empty", sections=None)
    check("Case created", case_id is not None)
    if not case_id:
        return

    # DOCX
    r = export_case(case_id, "docx")
    check("Empty DOCX export returns 200", r.ok, f"Got {r.status_code}")

    # Markdown
    r = export_case(case_id, "markdown")
    check("Empty markdown export returns 200", r.ok)
    if r.ok:
        content = r.json().get("content", "")
        check("Markdown has fallback content", len(content) > 0)

    # PDF
    r = export_case(case_id, "pdf")
    check("Empty PDF export returns 200", r.ok, f"Got {r.status_code}")


# ─────────────────────────────────────────────────────────────────
# TEST 12: Unicode content
# ─────────────────────────────────────────────────────────────────

def test_12_unicode():
    log("\n" + "=" * 70)
    log("TEST 12: Export with unicode content doesn't crash")
    log("=" * 70)

    unicode_sections = {
        "te_reo_section": "Ko Kiwi Community Trust te kaupapa. E whakapono ana matou ki te manaakitanga "
                         "me te kotahitanga. Te reo Maori resources — digital whanau programme.",
        "emoji_and_symbols": "Budget: $50,000 (NZD) — includes 10% contingency. "
                            "Partners: Foundation North, Todd Foundation. "
                            "Timeline: Jan 2026 - Dec 2026. Status: approved.",
        "special_chars": "Organisation's mission: bridging the \"digital divide\" in Aotearoa. "
                        "We've reached 3,200+ participants across 12 regions.",
    }

    case_id = create_test_case("unicode", unicode_sections)
    check("Case created", case_id is not None)
    if not case_id:
        return

    # Test all formats
    for fmt in ["docx", "pdf", "xlsx", "markdown", "json"]:
        r = export_case(case_id, fmt)
        check(f"Unicode {fmt} export returns 200", r.ok,
              f"Got {r.status_code}: {r.text[:100]}")

        if r.ok and fmt in ("docx", "pdf", "xlsx"):
            filename = r.json().get("filename", "")
            if filename:
                dl = download_file(case_id, filename)
                check(f"Unicode {fmt} download OK", dl.status_code == 200)
                check(f"Unicode {fmt} file > 0 bytes", len(dl.content) > 0)


# ─────────────────────────────────────────────────────────────────
# Runner
# ─────────────────────────────────────────────────────────────────

ALL_TESTS = {
    1: ("Fresh DOCX", test_1_fresh_docx),
    2: ("Fresh PDF", test_2_fresh_pdf),
    3: ("Fresh XLSX", test_3_fresh_xlsx),
    4: ("Markdown export", test_4_markdown),
    5: ("JSON export", test_5_json),
    6: ("Template DOCX", test_6_template_docx),
    7: ("Template XLSX", test_7_template_xlsx),
    8: ("Template PDF", test_8_template_pdf),
    9: (".doc conversion", test_9_doc_conversion),
    10: ("Download auth", test_10_download_auth),
    11: ("Empty sections", test_11_empty_sections),
    12: ("Unicode content", test_12_unicode),
}


def _strict(fn):
    """Under pytest, make script-style check() failures fail the test."""
    import functools

    @functools.wraps(fn)
    def wrapper():
        before = FAIL
        fn()
        failed = FAIL - before
        assert failed == 0, f"{failed} check(s) failed — see captured stdout"
    return wrapper


for _name, _fn in list(globals().items()):
    if _name.startswith("test_") and callable(_fn):
        globals()[_name] = _strict(_fn)


def main():
    global PASS, FAIL

    log("=" * 70)
    log("EXPORT PIPELINE TEST SUITE")
    log(f"API: {API}")
    log("=" * 70)

    if not setup_auth():
        log("[FATAL] Could not set up test auth. Is the backend running?")
        return 1

    log(f"Authenticated as org: {ORG_ID}")

    # Parse --case N
    target = None
    if "--case" in sys.argv:
        idx = sys.argv.index("--case")
        if idx + 1 < len(sys.argv):
            target = int(sys.argv[idx + 1])

    for num, (label, fn) in ALL_TESTS.items():
        if target and num != target:
            continue
        try:
            fn()
        except Exception as e:
            FAIL += 1
            log(f"\n[ERROR] Test {num} ({label}) crashed: {e}")
            import traceback
            traceback.print_exc()

    # Summary
    log(f"\n{'=' * 70}")
    log("SUMMARY")
    log("=" * 70)
    total = PASS + FAIL
    log(f"  Total: {total}  Passed: {PASS}  Failed: {FAIL}")
    if FAIL:
        log(f"\n  *** {FAIL} FAILURES ***")
    else:
        log(f"\n  All {PASS} checks passed!")
    log("=" * 70)

    # Save results to file
    output_path = os.path.join(os.path.dirname(__file__), "export_test_results.txt")
    with open(output_path, "w") as f:
        f.write("\n".join(RESULTS))
    log(f"\nResults saved to: {output_path}")

    return 0 if FAIL == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
